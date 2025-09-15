from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.units import inch
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph

class GeradorDeTermo:
    @staticmethod
    def gerar(formato, caminho_completo, dados_gerais, dados_agrupados):
        if formato == ".docx":
            GeradorDeTermo._gerar_docx(caminho_completo, dados_gerais, dados_agrupados)
        elif formato == ".pdf":
            GeradorDeTermo._gerar_pdf(caminho_completo, dados_gerais, dados_agrupados)
        else:
            raise NotImplementedError(f"Formato '{formato}' não suportado.")

    @staticmethod
    def _set_cell_style(cell, text, bold=False, size=10, font='Calibri'):
        cell.text = text
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    @staticmethod
    def _gerar_docx(caminho_completo, dados_gerais, dados_agrupados):
        doc = Document()
        section = doc.sections[0]
        section.orientation, section.width = section.height, section.width
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.add_run('Universidade Federal do Acre\n').bold = True
        p_titulo.add_run('Baixa de Bem Patrimonial').bold = True
        
        data_hora = datetime.now()
        p_data = doc.add_paragraph(f"Data: {data_hora.strftime('%d/%m/%Y')}\t\t\t\t\tHora: {data_hora.strftime('%H:%M')}")
        p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph(f"Termo: {dados_gerais.get('termo', 'N/A')}\t\t\t\t\tMotivo: {dados_gerais.get('motivo', 'N/A')}")
        doc.add_paragraph(f"Unidade Origem: {dados_gerais.get('unidade_origem', 'N/A')}")
        doc.add_paragraph(f"Unidade Destino: {dados_gerais.get('unidade_destino', 'N/A')}")
        doc.add_paragraph(f"Processo: {dados_gerais.get('processo', 'N/A')}")
        doc.add_paragraph()

        total_valor = 0.0
        colunas = ['Item', 'Registro', 'Plaqueta', 'Aceito', 'Descrição do Bem', 'Forma de Ingresso', 'Data Aquisição', 'Data Transferência', 'Valor']
        tabela = doc.add_table(rows=1, cols=len(colunas))
        tabela.style = 'Table Grid'
        hdr_cells = tabela.rows[0].cells
        for i, nome_coluna in enumerate(colunas):
            GeradorDeTermo._set_cell_style(hdr_cells[i], nome_coluna, bold=True)
        
        item_num = 1
        for (unidade, servidor), bens in dados_agrupados.items():
            for bem in bens:
                row_cells = tabela.add_row().cells
                
                # >>> CORREÇÃO AQUI: Acessando dados pelo nome da chave <<<
                tombo = bem.get('tombo', 'N/A')
                descricao = bem.get('descricao', 'N/A')
                data_aq = bem.get('data_aquisicao', 'N/A')
                valor_str = bem.get('valor', '0,00')
                forma_ingresso = bem.get('forma_ingresso', 'N/A')

                GeradorDeTermo._set_cell_style(row_cells[0], str(item_num))
                GeradorDeTermo._set_cell_style(row_cells[1], str(tombo))
                GeradorDeTermo._set_cell_style(row_cells[2], str(tombo))
                GeradorDeTermo._set_cell_style(row_cells[3], "S")
                GeradorDeTermo._set_cell_style(row_cells[4], descricao)
                GeradorDeTermo._set_cell_style(row_cells[5], forma_ingresso)
                GeradorDeTermo._set_cell_style(row_cells[6], str(data_aq))
                GeradorDeTermo._set_cell_style(row_cells[7], data_hora.strftime('%d/%m/%Y'))
                GeradorDeTermo._set_cell_style(row_cells[8], str(valor_str))
                
                try:
                    valor_float = float(str(valor_str).replace('.', '').replace(',', '.'))
                    total_valor += valor_float
                except (ValueError, TypeError): pass
                item_num += 1

        total_formatado = f"{total_valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        doc.add_paragraph()
        p_total = doc.add_paragraph(f"Total Página: {total_formatado}\t\t\t\t\tPágina: 1\nTotal Acumulado: {total_formatado}")
        p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save(caminho_completo)

    @staticmethod
    def _gerar_pdf(caminho_completo, dados_gerais, dados_agrupados):
        c = canvas.Canvas(caminho_completo, pagesize=landscape(letter))
        largura, altura = landscape(letter)
        margem = 0.5 * inch

        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(largura / 2, altura - margem, "Universidade Federal do Acre")
        c.drawCentredString(largura / 2, altura - margem - 14, "Baixa de Bem Patrimonial")
        data_hora = datetime.now()
        c.setFont("Helvetica", 9)
        c.drawRightString(largura - margem, altura - margem, f"Data: {data_hora.strftime('%d/%m/%Y')}")
        c.drawRightString(largura - margem, altura - margem - 12, f"Hora: {data_hora.strftime('%H:%M')}")
        
        y = altura - margem - 40
        c.setFont("Helvetica", 10)
        c.drawString(margem, y, f"Termo: {dados_gerais.get('termo', 'N/A')}")
        c.drawString(margem + 4 * inch, y, f"Motivo: {dados_gerais.get('motivo', 'N/A')}")
        y -= 15
        c.drawString(margem, y, f"Unidade Origem: {dados_gerais.get('unidade_origem', 'N/A')}")
        y -= 15
        c.drawString(margem, y, f"Unidade Destino: {dados_gerais.get('unidade_destino', 'N/A')}")
        y -= 15
        c.drawString(margem, y, f"Processo: {dados_gerais.get('processo', 'N/A')}")
        
        total_valor = 0.0
        styles = getSampleStyleSheet()
        style_normal = styles['Normal']
        style_normal.fontSize = 8
        style_normal.alignment = 1

        header = ['Item', 'Registro', 'Plaqueta', 'Aceito', 'Descrição do Bem', 'Forma de Ingresso', 'Data Aquisição', 'Data Transf.', 'Valor']
        dados_tabela = [header]
        item_num = 1
        for (unidade, servidor), bens in dados_agrupados.items():
            for bem in bens:
                # >>> CORREÇÃO AQUI: Acessando dados pelo nome da chave <<<
                tombo = bem.get('tombo', 'N/A')
                descricao = bem.get('descricao', 'N/A')
                data_aq = bem.get('data_aquisicao', 'N/A')
                valor_str = bem.get('valor', '0,00')
                forma_ingresso = bem.get('forma_ingresso', 'N/A')

                dados_tabela.append([
                    str(item_num), str(tombo), str(tombo), 'S',
                    Paragraph(descricao, style_normal), forma_ingresso, str(data_aq),
                    data_hora.strftime('%d/%m/%Y'), str(valor_str)
                ])
                try:
                    total_valor += float(str(valor_str).replace('.', '').replace(',', '.'))
                except (ValueError, TypeError): pass
                item_num += 1

        tabela = Table(dados_tabela, colWidths=[0.5*inch, 0.8*inch, 0.8*inch, 0.5*inch, 3.2*inch, 1*inch, 1*inch, 1*inch, 0.8*inch])
        tabela.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.black),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))

        y -= 25
        tabela.wrapOn(c, largura - 2 * margem, y)
        t_altura = tabela._height
        tabela.drawOn(c, margem, y - t_altura)

        y = y - t_altura - 20
        total_formatado = f"{total_valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        c.setFont("Helvetica", 10)
        c.drawRightString(largura - margem, y, f"Total Página: {total_formatado}")
        c.drawRightString(largura - margem, y - 15, f"Total Acumulado: {total_formatado}")
        
        c.save()